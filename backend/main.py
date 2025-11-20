import subprocess
import os
import sys
import time
import requests
import json
import urllib.parse
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
import sqlite3
import io
import base64
from contextlib import redirect_stdout
import traceback
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import r2_score, mean_squared_error
from scipy import stats
import warnings
from fastapi import Query, Depends
from sqlalchemy import create_engine, inspect

# Suppress warnings including bcrypt version warnings
warnings.filterwarnings('ignore')
import logging
logging.getLogger('passlib').setLevel(logging.ERROR)

# Only start Ollama if not already running
def is_ollama_running():
    import socket
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    try:
        s.connect(('localhost', 11434))
        s.close()
        return True
    except Exception:
        return False

if not is_ollama_running():
    # Start Ollama server in the background
    subprocess.Popen(['ollama', 'serve'], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

from fastapi import FastAPI, HTTPException, UploadFile, File, Query, Depends, status, Security
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
import requests
import random
import io
import os
from dotenv import load_dotenv
load_dotenv()

# For file processing
from typing import List
from PIL import Image
import pytesseract
import PyPDF2
import docx
from fastapi.responses import StreamingResponse
from fastapi.responses import JSONResponse
import json
from fastapi.responses import FileResponse

from sqlalchemy import create_engine, Column, Integer, String, ForeignKey, Text, DateTime
from sqlalchemy.orm import sessionmaker, declarative_base, relationship
from passlib.context import CryptContext
from jose import JWTError, jwt
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
import datetime
import shutil

# Determine base directory for data storage
if os.environ.get("DESKTOP_MODE") == "1":
    # In desktop mode, use the current working directory (which is set to user data dir)
    DATA_BASE_DIR = os.getcwd()
else:
    # In dev/server mode, use the directory of this file
    DATA_BASE_DIR = os.path.dirname(__file__)

UPLOAD_DIR = os.path.join(DATA_BASE_DIR, 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Create CSV upload directory
CSV_UPLOAD_DIR = os.path.join(DATA_BASE_DIR, 'csv_uploads')
os.makedirs(CSV_UPLOAD_DIR, exist_ok=True)

# --- User Auth & RBAC Setup ---
DATABASE_URL = "sqlite:///./users.db"
Base = declarative_base()
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True, nullable=False)
    password_hash = Column(String, nullable=False)
    role = Column(String, default="user")
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

class ChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    session_id = Column(String, nullable=False, index=True)
    prompt = Column(Text, nullable=False)
    response = Column(Text, nullable=False)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)
    user = relationship("User", backref="user_chat_history")

# New table for uploaded documents
class UploadedDocument(Base):
    __tablename__ = "uploaded_documents"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    session_id = Column(String, nullable=False, index=True)
    filename = Column(String, nullable=False)
    extracted_text = Column(Text, nullable=True)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)
    user = relationship("User", backref="user_uploaded_documents")

# New table for CSV files in data analysis mode
class UploadedCSV(Base):
    __tablename__ = "uploaded_csvs"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    session_id = Column(String, nullable=False, index=True)
    filename = Column(String, nullable=False)
    file_path = Column(String, nullable=False)
    columns_info = Column(Text, nullable=True)  # JSON string of column info
    table_name = Column(String, nullable=True)  # Name of the table where data is stored
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)
    user = relationship("User", backref="user_uploaded_csvs")

# Drop all tables and recreate them
def recreate_database():
    Base.metadata.drop_all(bind=engine)
    Base.metadata.create_all(bind=engine)

# Check if database schema is up to date and recreate if needed
def ensure_database_schema():
    try:
        # Try to query the uploaded_csvs table with table_name column
        from sqlalchemy import text
        with engine.connect() as conn:
            result = conn.execute(text("SELECT table_name FROM uploaded_csvs LIMIT 1"))
    except Exception as e:
        print(f"Database schema outdated, recreating tables: {e}")
        recreate_database()

# Create tables and ensure schema is up to date
Base.metadata.create_all(bind=engine)
ensure_database_schema()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = "supersecretkey"  # Change this in production
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 1 week

def safe_string_conversion(value, default=""):
    """Safely convert any value to string, handling lists and other types"""
    if isinstance(value, str):
        return value
    elif isinstance(value, list):
        return ' '.join(str(item) for item in value if item)
    elif value is None:
        return default
    else:
        return str(value)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: datetime.timedelta | None = None):
    to_encode = data.copy()
    expire = datetime.datetime.utcnow() + (expires_delta or datetime.timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/login")

# Dependency to get current user
from pydantic import BaseModel as PydanticBaseModel
class TokenData(PydanticBaseModel):
    username: str | None = None
    role: str | None = None

def get_current_user(token: str = Depends(oauth2_scheme), db=Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        role = payload.get("role")
        if not isinstance(username, str) or username is None:
            raise credentials_exception
        token_data = TokenData(username=username, role=role)
    except JWTError:
        raise credentials_exception
    user = db.query(User).filter(User.username == token_data.username).first()
    if user is None:
        raise credentials_exception
    return user

def require_role(required_role: str):
    def role_checker(user: User = Depends(get_current_user)):
        if getattr(user, 'role', None) != required_role:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return role_checker

# --- Registration Endpoint ---
class RegisterRequest(PydanticBaseModel):
    username: str
    password: str
    role: str = "user"  # Default role is 'user'

app = FastAPI()

# Add a set of allowed model names from ollama list
ALLOWED_OLLAMA_MODELS = {
    'mistral:latest',
    'anindya/prem1b-sql-ollama-fp116:latest',
    'llama3.1:latest',
    'qwen:0.5b',
    'gemma3:latest',
    'deepseek-v2:latest',
    'deepseek-coder:1.3b',
    'openchat:latest',
    'dolphin3:latest',
    'codellama:latest',
    'qwen2.5vl:latest',
    'deepseek-coder-v2:latest',
    'glm4:9b-chat-q4_0',
    'qwen3:0.6b',
    'llama3.2:1b',
    'deepseek-coder:latest',
    'llama3.1:8b',
    'nomic-embed-text:latest',
    'gpt-oss:latest'
}

class PromptInput(BaseModel):
    prompt: str
    model: str
    session_id: str  # Added session_id for context

class ChainRequest(BaseModel):
    prompt: str
    models: list[str]

# Data Analysis Models
class DataAnalysisRequest(BaseModel):
    prompt: str
    session_id: str
    model: str = "deepseek-coder-v2:latest"

class CSVUploadResponse(BaseModel):
    filename: str
    shape: tuple[int, int]  # Properly type the shape as a tuple of two integers
    columns: list[str]
    dtypes: dict[str, str]
    sample_data: list[dict]  # List of records from df.to_dict('records')

# Get API keys from environment variables
API_KEYS_STR = os.getenv("OPENROUTER_API_KEYS", "")
# Improved parsing to handle various formats
API_KEYS = []
if API_KEYS_STR:
    # Split by comma and handle various separators
    raw_keys = API_KEYS_STR.replace('\n', ',').replace(';', ',').split(',')
    for key in raw_keys:
        stripped_key = key.strip()
        if stripped_key and len(stripped_key) > 10:  # Basic validation for API key length
            API_KEYS.append(stripped_key)

print(f"ðŸ”‘ Loaded {len(API_KEYS)} API keys from environment variable")
if len(API_KEYS) < 50:  # If you expect 50 keys but got fewer
    print(f"âš ï¸  Warning: Expected 50 keys but only loaded {len(API_KEYS)}")
    print(f"   Raw environment variable length: {len(API_KEYS_STR)}")
    print(f"   Raw preview: {API_KEYS_STR[:200]}...")

# Get allowed origins from environment variables
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:3001,http://127.0.0.1:3000,http://127.0.0.1:3001,http://localhost:8080,http://127.0.0.1:8080")
ALLOWED_ORIGINS_LIST = [origin.strip() for origin in ALLOWED_ORIGINS.split(",")]
# Always ensure dev origins are present
for dev_origin in ["http://localhost:8080", "http://127.0.0.1:8080"]:
    if dev_origin not in ALLOWED_ORIGINS_LIST:
        ALLOWED_ORIGINS_LIST.append(dev_origin)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS_LIST,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_URL = "https://openrouter.ai/api/v1/chat/completions"

# ============================================================================
# DATA ANALYSIS FUNCTIONS (Converted from Streamlit)
# ============================================================================

# Ollama configuration for data analysis
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_URL = f"{OLLAMA_BASE_URL}/api/generate"
DATA_ANALYSIS_MODEL = "deepseek-coder-v2:latest"
MAX_RETRIES = 3
REQUEST_TIMEOUT = 300

def query_ollama_data_analysis(prompt, model=DATA_ANALYSIS_MODEL, timeout=REQUEST_TIMEOUT, retries=MAX_RETRIES):
    """Query Ollama API for data analysis with better error handling"""
    url = f"{OLLAMA_BASE_URL}/api/generate"
    headers = {"Content-Type": "application/json"}
    data = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1,
            "top_p": 0.9,
            "num_ctx": 4096
        }
    }
    
    for attempt in range(retries):
        try:
            response = requests.post(
                url,
                headers=headers,
                json=data,
                timeout=timeout
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'response' in result:
                    return result['response']
                else:
                    return "Error: Unexpected response format from Ollama"
            else:
                error_msg = f"HTTP {response.status_code}"
                try:
                    error_data = response.json()
                    if 'error' in error_data:
                        error_msg = error_data['error']
                except:
                    error_msg = response.text or error_msg
                
                if attempt == retries - 1:  # Last attempt
                    return f"API Error: {error_msg}"
                
        except requests.exceptions.Timeout:
            if attempt == retries - 1:  # Last attempt
                return "TIMEOUT_ERROR"
                
        except Exception as e:
            if attempt == retries - 1:  # Last attempt
                return f"Connection Error: {str(e)}"
        
        # Exponential backoff before retry
        if attempt < retries - 1:
            wait_time = (2 ** attempt) * 2  # 2, 4, 8 seconds
            time.sleep(wait_time)
    
    return "Failed after multiple attempts"

def generate_fallback_code(question, df):
    """
    Intelligent fallback code generator.
    Detects question intent (statistics, correlation, bar chart, histogram, scatter plot, etc.)
    and generates suitable Python code automatically.
    """
    import re
    question_lower = question.lower()
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()

    # --- 1ï¸âƒ£ Detect if user asks for summary stats ---
    if any(k in question_lower for k in ["average", "mean", "minimum", "maximum", "max", "min", "median", "summary", "describe"]):
        for col in numeric_cols:
            if col.lower() in question_lower:
                return f"""
# Summary stats for {col}
avg = df["{col}"].mean()
min_val = df["{col}"].min()
max_val = df["{col}"].max()
median_val = df["{col}"].median()
print("Column: {col}")
print(f"Average (Mean): {{avg:.2f}}")
print(f"Median: {{median_val:.2f}}")
print(f"Minimum: {{min_val}}")
print(f"Maximum: {{max_val}}")
"""
        # If no specific column found
        return """
print("Summary Statistics for Numeric Columns:")
print(df.describe())
"""

    # --- 2ï¸âƒ£ Detect histogram/distribution questions ---
    if any(k in question_lower for k in ["distribution", "histogram", "frequency", "spread"]):
        target_col = next((col for col in numeric_cols if col.lower() in question_lower), numeric_cols[0])
        return f"""
# Histogram for {target_col}
import matplotlib.pyplot as plt
plt.figure(figsize=(8,5))
plt.hist(df["{target_col}"].dropna(), bins=20, color='skyblue', edgecolor='black')
plt.title("Distribution of {target_col}")
plt.xlabel("{target_col}")
plt.ylabel("Frequency")
plt.tight_layout()
plt.show()
"""

    # --- 3ï¸âƒ£ Detect correlation or "relationship" type questions ---
    if any(k in question_lower for k in ["correlation", "relationship", "compare", "association"]):
        if "heatmap" in question_lower or "matrix" in question_lower:
            return """
# Correlation heatmap
import matplotlib.pyplot as plt
import seaborn as sns
numeric_df = df.select_dtypes(include=['number'])
plt.figure(figsize=(10,8))
sns.heatmap(numeric_df.corr(), annot=True, cmap='coolwarm', fmt='.2f')
plt.title('Correlation Heatmap')
plt.tight_layout()
plt.show()
print(numeric_df.corr())
"""
        # Otherwise, scatter plot for "A vs B"
        matches = re.findall(r"([a-z_]+)\s*(?:vs|against|and)\s*([a-z_]+)", question_lower)
        if matches:
            col1, col2 = matches[0]
            col1 = next((c for c in df.columns if col1 in c.lower()), numeric_cols[0])
            col2 = next((c for c in df.columns if col2 in c.lower()), numeric_cols[1])
            return f"""
# Scatter plot between {col1} and {col2}
import matplotlib.pyplot as plt
plt.figure(figsize=(8,6))
plt.scatter(df["{col1}"], df["{col2}"], alpha=0.6, color='teal')
plt.title("{col1} vs {col2}")
plt.xlabel("{col1}")
plt.ylabel("{col2}")
plt.tight_layout()
plt.show()
"""

    # --- 4ï¸âƒ£ Detect "by" or grouped average (bar chart) ---
    if "by" in question_lower or "group" in question_lower or "range" in question_lower:
        target_col = next((col for col in numeric_cols if col.lower() in question_lower and "by" not in col.lower()), "median_income")
        group_col = None
        for col in df.columns:
            if col.lower() in question_lower.split("by")[-1]:
                group_col = col
                break
        if not group_col:
            group_col = numeric_cols[1] if len(numeric_cols) > 1 else df.columns[0]

        return f"""
# Grouped bar chart of average {target_col} by {group_col} range
import matplotlib.pyplot as plt
import pandas as pd
df['{group_col}_range'] = pd.cut(df['{group_col}'], bins=5)
avg_data = df.groupby('{group_col}_range')['{target_col}'].mean()

plt.figure(figsize=(10,6))
avg_data.plot(kind='bar', color='cornflowerblue', edgecolor='black')
plt.title('Average {target_col} by {group_col} Range')
plt.xlabel('{group_col} Range')
plt.ylabel('Average {target_col}')
plt.xticks(rotation=45)
plt.tight_layout()
plt.show()
print(avg_data)
"""

    # --- 5ï¸âƒ£ Detect trend / time-based questions ---
    if any(k in question_lower for k in ["trend", "over time", "year", "month", "timeline", "progression"]):
        target_col = next((col for col in numeric_cols if col.lower() in question_lower), numeric_cols[0])
        return f"""
# Line chart showing trend of {target_col} over index
import matplotlib.pyplot as plt
plt.figure(figsize=(8,5))
plt.plot(df.index, df["{target_col}"], color='orange')
plt.title("Trend of {target_col} over dataset index")
plt.xlabel("Index")
plt.ylabel("{target_col}")
plt.tight_layout()
plt.show()
"""

    # --- Default generic analysis ---
    return """
print("Dataset Shape:", df.shape)
print("\\nColumn Info:")
print(df.dtypes)
print("\\nBasic Statistics:")
print(df.describe())
print("\\nMissing Values:")
print(df.isnull().sum())
"""


def execute_data_analysis_code(code, df, table_name=None, globals_dict=None):
    """Safely execute Python code with dataframe context and SQL support"""
    if globals_dict is None:
        globals_dict = {}
    
    # Create SQLite engine if table_name is provided
    engine = None
    if table_name:
        engine = create_engine(f'sqlite:///databases/analysis.db')
    
    # Prepare execution environment with the actual dataframe
    exec_globals = {
        'df': df,
        'pd': pd,
        'np': np,
        'plt': plt,
        'sns': sns,
        'px': px,
        'go': go,
        'stats': stats,
        'LinearRegression': LinearRegression,
        'train_test_split': train_test_split,
        'r2_score': r2_score,
        'mean_squared_error': mean_squared_error,
        'execute_sql': lambda query: execute_sql_query(query, engine) if engine else None,
        'table_name': table_name,
        **globals_dict
    }
    
    # Clean the code - remove any file loading attempts
    cleaned_code = code.replace("pd.read_csv('path_to_your_dataset.csv')", "df")
    cleaned_code = cleaned_code.replace("pd.read_csv(", "# pd.read_csv(")
    cleaned_code = cleaned_code.replace("df = pd.read_csv", "# df = pd.read_csv")
    
    # Capture output
    output = io.StringIO()
    chart_data = None
    
    try:
        with redirect_stdout(output):
            exec(cleaned_code, exec_globals)
        
        # Get printed output
        printed_output = output.getvalue()
        
        # Check if a plot was created and save it as base64
        if plt.get_fignums():
            # Get the current figure
            fig = plt.gcf()
            
            # Save plot to base64 string
            buffer = io.BytesIO()
            fig.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            plot_data = buffer.getvalue()
            buffer.close()
            
            # Convert to base64
            chart_data = base64.b64encode(plot_data).decode('utf-8')
            
            # Clear the figure
            plt.clf()
            plt.close('all')
        
        return {
            'success': True,
            'output': printed_output,
            'chart': chart_data,
            'globals': exec_globals
        }
    
    except Exception as e:
        plt.close('all')  # Clean up any plots
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }

def save_to_sql_database(df, table_name, db_path="databases/analysis.db"):
    """Save DataFrame to SQLite database"""
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    
    # Create SQLite connection
    engine = create_engine(f'sqlite:///{db_path}')
    
    # Save DataFrame to SQL
    df.to_sql(table_name, engine, if_exists='replace', index=False)
    return engine

def execute_sql_query(query, engine):
    """Execute SQL query, print results, and return DataFrame"""
    try:
        df = pd.read_sql_query(query, engine)
        # Print a readable summary of the results
        print("SQL Query Results:")
        print(df.head(20).to_string(index=False))
        return df
    except Exception as e:
        print(f"SQL Query Error: {str(e)}")
        raise Exception(f"SQL Query Error: {str(e)}")
    
def get_database_engine():
    db_path = "databases/analysis.db"
    return create_engine(f"sqlite:///{db_path}")

@app.get("/list-tables")
def list_uploaded_tables(current_user: User = Depends(get_current_user)):
    engine = get_database_engine()
    inspector = inspect(engine)
    tables = inspector.get_table_names()
    if not tables:
        raise HTTPException(status_code=404, detail="No tables found in the database.")
    return {"tables": tables}

@app.get("/describe-table")
def describe_table(
    table_name: str = Query(..., description="Name of the table to describe"),
    current_user: User = Depends(get_current_user)
):
    engine = get_database_engine()
    inspector = inspect(engine)
    columns = inspector.get_columns(table_name)
    if not columns:
        raise HTTPException(status_code=404, detail=f"Table '{table_name}' not found.")
    return {
        "table": table_name,
        "columns": [
            {"name": col["name"], "type": str(col["type"]), "nullable": col["nullable"]}
            for col in columns
        ]
    }
@app.post("/execute-sql")
def execute_sql(
    query: str = Query(..., description="SQL query to execute (SELECT only)"),
    current_user: User = Depends(get_current_user)
):
    if not query.strip().lower().startswith("select"):
        raise HTTPException(status_code=400, detail="Only SELECT queries are allowed for safety.")

    engine = get_database_engine()
    try:
        df = pd.read_sql_query(query, engine)
        records = df.head(50).to_dict(orient="records")  # Limit rows for performance
        return {"success": True, "row_count": len(df), "columns": list(df.columns), "data": records}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"SQL execution error: {str(e)}")
    

def load_csv_file(file_path):
    """Load CSV file with multiple fallback methods"""
    import chardet
    
    # Read the file content first to detect encoding
    with open(file_path, 'rb') as f:
        raw_content = f.read()
    
    # Try to detect encoding
    try:
        detected_enc = chardet.detect(raw_content)['encoding']
        print(f"Detected encoding: {detected_enc}")
        if detected_enc is None:
            detected_enc = 'utf-8'
    except Exception as e:
        print(f"Error detecting encoding: {str(e)}")
        detected_enc = 'utf-8'
    
    # Try different methods to read the file
    methods = [
        (lambda: pd.read_csv(file_path, encoding=detected_enc), f"detected encoding ({detected_enc})"),
        (lambda: pd.read_csv(file_path, encoding='utf-8'), "utf-8"),
        (lambda: pd.read_csv(file_path, encoding='latin1'), "latin1"),
        (lambda: pd.read_csv(file_path, encoding='cp1252'), "cp1252"),
        (lambda: pd.read_csv(file_path, encoding='iso-8859-1'), "iso-8859-1"),
        (lambda: pd.read_csv(file_path, sep=';', encoding=detected_enc), f"semicolon separator, {detected_enc}"),
        (lambda: pd.read_csv(file_path, sep='\t', encoding=detected_enc), f"tab separator, {detected_enc}"),
        # Additional fallback methods for problematic files
        (lambda: pd.read_csv(file_path, encoding=detected_enc, on_bad_lines='skip'), f"skip bad lines, {detected_enc}"),
        (lambda: pd.read_csv(file_path, encoding=detected_enc, engine='python'), f"python engine, {detected_enc}"),
        (lambda: pd.read_csv(file_path, encoding=detected_enc, delimiter=None), f"auto delimiter, {detected_enc}")
    ]
    
    errors = []
    for method, description in methods:
        try:
            print(f"Trying to read CSV with {description}")
            df = method()
            if df is not None and not df.empty:
                print(f"Successfully read CSV with {description}")
                print(f"DataFrame shape: {df.shape}")
                print(f"Columns: {df.columns.tolist()}")
                
                # Clean column names
                df.columns = df.columns.str.strip()
                for col in df.columns:
                    try:
                        df[col] = df[col].replace(',', '', regex=True).astype(float)
                    except Exception:
                        pass  # ignore non-numeric columns

                print(f"Converted dtypes:\n{df.dtypes}")
                return df

        except Exception as e:
            error_msg = f"Method '{description}' failed: {str(e)}"
            print(error_msg)
            errors.append(error_msg)
            continue
    
    # If we get here, all methods failed
    error_details = "\n".join(errors)
    print(f"All methods failed to read the CSV file. Errors:\n{error_details}")
    return None

# ============================================================================
# EXISTING ENDPOINTS
# ============================================================================

# âœ… Health check endpoint
@app.get("/health")
def health_check():
    return {"status": "healthy", "message": "Backend is running"}

# âœ… API Keys diagnostic endpoint
@app.get("/debug/keys")
def debug_api_keys():
    """Diagnostic endpoint to check API key parsing"""
    raw_keys_str = os.getenv("OPENROUTER_API_KEYS", "")
    parsed_keys = [key.strip() for key in raw_keys_str.split(",") if key.strip()]
    
    return {
        "raw_environment_variable_length": len(raw_keys_str),
        "raw_environment_variable_preview": raw_keys_str[:100] + "..." if len(raw_keys_str) > 100 else raw_keys_str,
        "total_keys_after_parsing": len(parsed_keys),
        "parsed_keys_preview": [key[:10] + "..." + key[-4:] if len(key) > 14 else key for key in parsed_keys[:5]],
        "all_parsed_keys_count": len(parsed_keys),
        "environment_variable_name": "OPENROUTER_API_KEYS"
    }

# âœ… Main endpoint to send prompt and receive model response
@app.post("/ask")
def ask_model(data: PromptInput, db=Depends(get_db), current_user: User = Depends(get_current_user)):
    # Fetch last 1 message for context (oldest first)
    history = db.query(ChatHistory)\
        .filter(ChatHistory.user_id == current_user.id, ChatHistory.session_id == data.session_id)\
        .order_by(ChatHistory.timestamp.desc())\
        .limit(1).all()

    context = ""
    for msg in reversed(history):
        context += f"User: {msg.prompt}\nAI: {msg.response}\n"

    # Fetch all extracted PDF texts for this session and user (in upload order)
    docs = db.query(UploadedDocument).filter(UploadedDocument.user_id == current_user.id, UploadedDocument.session_id == data.session_id, UploadedDocument.extracted_text != None).order_by(UploadedDocument.timestamp.asc()).all()

    pdf_context = ""
    for idx, doc in enumerate(docs):
        if doc.extracted_text:
            pdf_context += f"[PDF {idx+1} Content Start: {doc.filename}]\n{doc.extracted_text}\n[PDF {idx+1} Content End]\n"

    # NEW: Search SearxNG for relevant information
    searxng_context = ""
    try:
        searxng_url = "http://localhost:8888/search"  # Using local SearxNG instance
        params = {
            "q": data.prompt,
            "format": "json",
            "categories": "general",
            "engines": "google,bing,duckduckgo",
            "pageno": 1
        }
        
        search_response = requests.get(searxng_url, params=params, timeout=10)
        if search_response.ok:
            search_data = search_response.json()
            search_results = search_data.get("results", [])[:5]  # Limit to 5 results
            
            if search_results:
                searxng_context = "\n[SEARCH RESULTS START]\n"
                for idx, result in enumerate(search_results, 1):
                    title = result.get("title", "No title")
                    content = result.get("content", "No content")
                    url = result.get("url", "No URL")
                    searxng_context += f"Source {idx}: {title}\nURL: {url}\nContent: {content}\n\n"
                searxng_context += "[SEARCH RESULTS END]\n\n"
    except Exception as e:
        print(f"SearxNG search failed: {e}")

    # Combine all contexts
    full_prompt = ""
    if pdf_context:
        full_prompt += pdf_context
    if searxng_context:
        full_prompt += searxng_context
    
    full_prompt += context + f"Based on the above search results and context, please answer the following question: {data.prompt}\nAI:"

    # Use Ollama local API
    ollama_url = "http://localhost:11434/api/generate"
    payload = {
        "model": data.model,
        "prompt": full_prompt,
        "stream": True,
        "num_predict": 256,  # Increased for more detailed responses
        "temperature": 0.7,
    }

    if data.model not in ALLOWED_OLLAMA_MODELS:
        raise HTTPException(status_code=400, detail=f"Model '{data.model}' is not available on this server.")

    def stream_ollama():
        # First, yield the search results as a separate message
        if searxng_context:
            yield json.dumps({
                "type": "search_results", 
                "content": searxng_context,
                "done": False
            }) + "\n"
        
        try:
            with requests.post(ollama_url, json=payload, stream=True, timeout=600) as response:
                response.raise_for_status()
                for line in response.iter_lines():
                    if line:
                        try:
                            chunk = line.decode()
                            data_json = json.loads(chunk)
                            resp = data_json.get("response", "")
                            done = data_json.get("done", False)
                            yield json.dumps({
                                "type": "model_response",
                                "response": resp, 
                                "done": done
                            }) + "\n"
                            if done:
                                break
                        except Exception:
                            continue
            # Ensure the last chunk is sent
            yield json.dumps({"type": "model_response", "done": True}) + "\n"
        except Exception as e:
            yield json.dumps({
                "type": "error",
                "error": f"Ollama error: {str(e)}", 
                "done": True
            }) + "\n"

    return StreamingResponse(stream_ollama(), media_type="application/jsonl")

# ============================================================================
# NEW DATA ANALYSIS ENDPOINTS
# ============================================================================

# Ensure upload directories exist - use backend directory for CSV uploads
CSV_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "csv_uploads")
os.makedirs(CSV_UPLOAD_DIR, exist_ok=True)

from fastapi.middleware.cors import CORSMiddleware

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/upload-csv")
async def upload_csv(
    file: UploadFile = File(...),
    session_id: str = Query(..., min_length=1, description="Session ID for tracking uploads"),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Upload CSV file for data analysis"""
    print(f"Received upload request for session: {session_id}")
    
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
        
    if not file.filename.lower().endswith('.csv'):
        raise HTTPException(
            status_code=400, 
            detail=f"Invalid file type: {file.filename}. Only CSV files are allowed"
        )
    
    # Create unique filename to avoid conflicts
    timestamp = int(time.time())
    filename = f"{current_user.id}_{session_id}_{timestamp}_{file.filename}"
    file_path = os.path.join(CSV_UPLOAD_DIR, filename)
    
    print(f"Processing file: {filename}")
    print(f"File size: {file.size if hasattr(file, 'size') else 'unknown'}")
    print(f"Content type: {file.content_type}")
    
    # Save file
    try:
        print("Reading file content...")
        content = await file.read()
        if not content:
            error_msg = "Uploaded file is empty"
            print(f"Error: {error_msg}")
            raise HTTPException(status_code=400, detail=error_msg)
        
        print(f"Read {len(content)} bytes, writing to {file_path}")
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        # Verify file was written
        if not os.path.exists(file_path):
            error_msg = f"File was not saved to {file_path}"
            print(f"Error: {error_msg}")
            raise HTTPException(status_code=500, detail=error_msg)
            
        file_size = os.path.getsize(file_path)
        print(f"File saved successfully. Size: {file_size} bytes")
        
        if file_size == 0:
            error_msg = "Saved file is empty"
            print(f"Error: {error_msg}")
            os.remove(file_path)  # Clean up empty file
            raise HTTPException(status_code=400, detail=error_msg)
            
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    
    # Try to load and analyze the CSV
    try:
        print(f"Attempting to load CSV file: {file_path}")
        df = load_csv_file(file_path)
        if df is None or df.empty:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(
                status_code=400, 
                detail="Could not read CSV file. The file might have an unsupported format or encoding. Please ensure it's a valid CSV file with UTF-8 or similar encoding."
            )
        
        # Get basic info about the dataset
        try:
            sample_data = df.head(5).to_dict('records') if len(df) > 0 else {}
            
            columns_info = {
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.astype(str).to_dict(),
                'shape': df.shape,
                'sample_data': sample_data
            }
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=400, detail=f"Error processing CSV data: {str(e)}")
        
        # Store the CSV data in SQLite
        table_name = f"data_{current_user.id}_{session_id}_{int(time.time())}"
        engine = save_to_sql_database(df, table_name)
        
        # Store metadata in database
        csv_record = UploadedCSV(
            user_id=current_user.id,
            session_id=session_id,
            filename=file.filename,
            file_path=file_path,
            columns_info=json.dumps(columns_info),
            table_name=table_name
        )
        db.add(csv_record)
        db.commit()
        db.refresh(csv_record)
        
        # Create response using the Pydantic model
        response_data = CSVUploadResponse(
            filename=file.filename,
            shape=df.shape,
            columns=df.columns.tolist(),
            dtypes=df.dtypes.astype(str).to_dict(),
            sample_data=sample_data
        )
        
        return response_data
        
    except Exception as e:
        # Clean up file if processing failed
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=400, detail=f"Error processing CSV: {str(e)}")

@app.post("/data-analysis")
def data_analysis(
    data: DataAnalysisRequest,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Process data analysis questions"""
    
    # Get the uploaded CSV for this session
    csv_record = db.query(UploadedCSV).filter(
        UploadedCSV.user_id == current_user.id,
        UploadedCSV.session_id == data.session_id
    ).order_by(UploadedCSV.timestamp.desc()).first()
    
    if not csv_record:
        raise HTTPException(status_code=400, detail="No CSV file uploaded for this session")
    
    if not os.path.exists(csv_record.file_path):
        raise HTTPException(status_code=400, detail="CSV file not found. Please re-upload.")
    
    # Load the CSV
    try:
        df = load_csv_file(csv_record.file_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not load CSV file")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error loading CSV: {str(e)}")
    
    def stream_data_analysis():
        try:
            # Check if Ollama is available
            ollama_available = True
            try:
                response = requests.get(f"{OLLAMA_BASE_URL}/api/version", timeout=5)
                if response.status_code != 200:
                    ollama_available = False
            except:
                ollama_available = False
            
            if ollama_available and data.model in ALLOWED_OLLAMA_MODELS:
                # Generate code using AI
                yield json.dumps({
                    "type": "status",
                    "message": "Generating analysis code with AI...",
                    "done": False
                }) + "\n"
                
                context = f"""
Dataset: {df.shape[0]} rows, {df.shape[1]} columns
Columns: {list(df.columns)}
Types: {dict(df.dtypes)}

RULES:
1. Use 'df' (already loaded)
2. Generate executable Python code only
3. Use print() for results
4. Keep code concise
5. Include matplotlib plots when appropriate

Question: {data.prompt}
Code:"""
                
                ai_response = query_ollama_data_analysis(context, model=data.model)
                
                if "Error" in ai_response or "TIMEOUT" in ai_response:
                    yield json.dumps({
                        "type": "status",
                        "message": "AI unavailable, using fallback code generation...",
                        "done": False
                    }) + "\n"
                    code = generate_fallback_code(data.prompt, df)
                else:
                    # Extract code from AI response
                    if '```python' in ai_response:
                        code_start = ai_response.find('```python') + 9
                        code_end = ai_response.find('```', code_start)
                        code = ai_response[code_start:code_end].strip()
                    else:
                        code = ai_response.strip()
            else:
                yield json.dumps({
                    "type": "status",
                    "message": "Using fallback code generation...",
                    "done": False
                }) + "\n"
                code = generate_fallback_code(data.prompt, df)
            
            # Send the generated code
            yield json.dumps({
                "type": "code",
                "content": code,
                "done": False
            }) + "\n"
            
            # Execute the code
            yield json.dumps({
                "type": "status",
                "message": "Executing analysis...",
                "done": False
            }) + "\n"
            
            result = execute_data_analysis_code(code, df, csv_record.table_name)
            
            if result['success']:
                # Send output if any
                if result.get('output'):
                    yield json.dumps({
                        "type": "output",
                        "content": result['output'],
                        "done": False
                    }) + "\n"
                
                # Send chart if any
                if result.get('chart'):
                    yield json.dumps({
                        "type": "chart",
                        "content": result['chart'],
                        "done": False
                    }) + "\n"
                
                # Send explanation
                explanation = f"Analysis completed successfully. The results show patterns and insights from your dataset with {df.shape} rows and {df.shape} columns."
                yield json.dumps({
                    "type": "explanation",
                    "content": explanation,
                    "done": False
                }) + "\n"
                
            else:
                yield json.dumps({
                    "type": "error",
                    "content": f"Execution error: {result['error']}",
                    "done": False
                }) + "\n"
            
            # Final done signal
            yield json.dumps({
                "type": "analysis_complete",
                "done": True
            }) + "\n"
            
        except Exception as e:
            yield json.dumps({
                "type": "error",
                "content": f"Analysis error: {str(e)}",
                "done": True
            }) + "\n"
    
    return StreamingResponse(stream_data_analysis(), media_type="application/jsonl")

@app.get("/csv-info")
def get_csv_info(
    session_id: str = Query(...),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get information about uploaded CSV for this session"""
    csv_record = db.query(UploadedCSV).filter(
        UploadedCSV.user_id == current_user.id,
        UploadedCSV.session_id == session_id
    ).order_by(UploadedCSV.timestamp.desc()).first()
    
    if not csv_record:
        return {"has_csv": False}
    
    columns_info = json.loads(csv_record.columns_info) if csv_record.columns_info else {}
    
    return {
        "has_csv": True,
        "filename": csv_record.filename,
        "uploaded_at": csv_record.timestamp.isoformat(),
        **columns_info
    }

# ============================================================================
# EXISTING ENDPOINTS (CONTINUED)
# ============================================================================

@app.get("/search/searxng")
def searxng_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search using local SearxNG instance."""
    searxng_url = "http://localhost:8888/search"  # Using local SearxNG instance
    
    params = {
        "q": query,
        "format": "json",
        "categories": "general",
        "engines": "google,bing,duckduckgo",
        "pageno": 1
    }
    
    try:
        response = requests.get(searxng_url, params=params, timeout=15)
        response.raise_for_status()
        
        data = response.json()
        
        # Extract and format results
        results = []
        for result in data.get("results", [])[:max_results]:
            results.append({
                "title": result.get("title", ""),
                "content": result.get("content", ""),
                "url": result.get("url", ""),
                "engine": result.get("engine", "")
            })
        
        return {"results": results, "query": query}
        
    except Exception as e:
        return {"error": str(e), "results": []}

@app.post("/chain")
def chain_models(data: ChainRequest):
    def stream_chain():
        current_prompt = data.prompt
        for model_id in data.models:
            if model_id not in ALLOWED_OLLAMA_MODELS:
                yield json.dumps({"model": model_id, "response": f"Model '{model_id}' is not available on this server.", "done": True}) + "\n"
                current_prompt = data.prompt
                continue
            ollama_url = "http://localhost:11434/api/generate"
            payload = {
                "model": model_id,
                "prompt": current_prompt,
                "stream": True,
                "num_predict": 64,
                "temperature": 0.8,
            }
            full_response = ""
            try:
                with requests.post(ollama_url, json=payload, stream=True, timeout=600) as response:
                    response.raise_for_status()
                    for line in response.iter_lines():
                        if line:
                            try:
                                chunk = line.decode()
                                data_json = json.loads(chunk)
                                resp = data_json.get("response", "")
                                done = data_json.get("done", False)
                                full_response += resp
                                yield json.dumps({"model": model_id, "response": resp, "done": False}) + "\n"
                                if done:
                                    break
                            except Exception:
                                continue
                    # After model is done, send done for this model
                    yield json.dumps({"model": model_id, "done": True}) + "\n"
                    # Pass the full output to the next model
                    current_prompt = full_response
            except Exception as e:
                yield json.dumps({"model": model_id, "response": f"Ollama error: {str(e)}", "done": True}) + "\n"
                current_prompt = data.prompt
    return StreamingResponse(stream_chain(), media_type="application/jsonl")

@app.post("/upload")
def upload_files(files: List[UploadFile] = File(...), session_id: str = Query(...), db=Depends(get_db), current_user: User = Depends(get_current_user)):
    saved_files = []
    for file in files:
        if file.filename is None:
            continue
        filename = file.filename
        file_path = os.path.join(UPLOAD_DIR, filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        extracted_text = None
        
        # Handle file types
        if filename.lower().endswith(".pdf"):
            try:
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    extracted_text = text
            except Exception as e:
                extracted_text = None
        elif filename.lower().endswith(".docx"):
            try:
                docx_file = docx.Document(file_path)
                text = "\n".join([para.text for para in docx_file.paragraphs])
                extracted_text = text
            except Exception as e:
                extracted_text = None
        elif filename.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")):
            try:
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image)
                extracted_text = text
            except Exception as e:
                extracted_text = None
        
        # Store in UploadedDocument
        doc = UploadedDocument(
            user_id=current_user.id,
            session_id=session_id,
            filename=filename,
            extracted_text=extracted_text
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        saved_files.append({"filename": filename, "extracted_text": bool(extracted_text)})
    return {"files": saved_files}

# Continue with remaining endpoints...
EXA_API_KEY = os.getenv("EXA_API_KEY", "")

@app.get("/search/web")
def exa_web_search(query: str = Query(..., description="Search query")):
    if not EXA_API_KEY:
        return {"error": "EXA_API_KEY not set in environment variables."}
    url = "https://api.exa.ai/search"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "Authorization": f"Bearer {EXA_API_KEY}"
    }
    payload = {
        "query": query,
        "numResults": 5
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "")

@app.get("/search/youtube")
def tavily_youtube_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search YouTube videos using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "youtube"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily YouTube response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

@app.get("/search/reddit")
def tavily_reddit_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search Reddit posts using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "reddit"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily Reddit response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

@app.get("/search/academic")
def tavily_academic_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search academic papers using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "academic"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily Academic response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

@app.get("/search/crypto")
def tavily_crypto_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search crypto news using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "crypto"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily Crypto response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

# --- Registration Endpoint ---
@app.post("/register")
def register_user(data: RegisterRequest, db=Depends(get_db)):
    if db.query(User).filter(User.username == data.username).first():
        raise HTTPException(status_code=400, detail="Username already registered")
    user = User(
        username=data.username,
        password_hash=get_password_hash(data.password),
        role=data.role,
        created_at=datetime.datetime.utcnow()  # Set creation time
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return {"msg": "User registered successfully"}

# --- Login Endpoint ---
class Token(PydanticBaseModel):
    access_token: str
    token_type: str

@app.post("/login", response_model=Token)
def login(form_data: OAuth2PasswordRequestForm = Depends(), db=Depends(get_db)):
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.password_hash):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    access_token = create_access_token(data={"sub": user.username, "role": user.role})
    return {"access_token": access_token, "token_type": "bearer"}

# --- Get Current User ---
@app.get("/me")
def read_users_me(current_user: User = Depends(get_current_user)):
    return {"username": current_user.username, "role": current_user.role}

# Example of RBAC-protected endpoint
@app.get("/admin-only")
def admin_only_endpoint(user: User = Depends(require_role("admin"))):
    return {"msg": f"Hello, admin {user.username}!"}

# --- Chat History ---
from fastapi import Body

class ChatMessageRequest(PydanticBaseModel):
    session_id: str
    prompt: str
    response: str

class ChatMessageResponse(PydanticBaseModel):
    id: int
    session_id: str
    prompt: str
    response: str
    timestamp: datetime.datetime

from typing import List as TypingList

@app.post("/history", response_model=ChatMessageResponse)
def save_chat_message(
    data: ChatMessageRequest,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chat = ChatHistory(
        user_id=current_user.id,
        session_id=data.session_id,
        prompt=data.prompt,
        response=data.response
    )
    db.add(chat)
    db.commit()
    db.refresh(chat)
    return ChatMessageResponse(
        id=getattr(chat, 'id'),
        session_id=getattr(chat, 'session_id'),
        prompt=getattr(chat, 'prompt'),
        response=getattr(chat, 'response'),
        timestamp=getattr(chat, 'timestamp')
    )

from fastapi import Query as FastAPIQuery

@app.get("/history", response_model=TypingList[ChatMessageResponse])
def get_chat_history(
    session_id: str = FastAPIQuery(..., description="Session ID to filter chat history"),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chats = db.query(ChatHistory).filter(ChatHistory.user_id == current_user.id, ChatHistory.session_id == session_id).order_by(ChatHistory.timestamp.asc()).all()
    return [
        ChatMessageResponse(
            id=getattr(chat, 'id'),
            session_id=getattr(chat, 'session_id'),
            prompt=getattr(chat, 'prompt'),
            response=getattr(chat, 'response'),
            timestamp=getattr(chat, 'timestamp')
        ) for chat in chats
    ]

# Serve React build files (dist) as static files
if getattr(sys, 'frozen', False):
    # Running in a PyInstaller bundle
    frontend_build_dir = os.path.join(sys._MEIPASS, 'dist')
else:
    # Running in normal Python
    frontend_build_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'dist'))

print("Serving static files from:", frontend_build_dir)
print("index.html exists:", os.path.exists(os.path.join(frontend_build_dir, "index.html")))

from starlette.staticfiles import StaticFiles
from starlette.responses import FileResponse
from starlette.requests import Request

class SPAStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        print(f"Requested path: {path}")
        full_path = os.path.join(self.directory, path)
        if not os.path.exists(full_path) or os.path.isdir(full_path):
            index_path = os.path.join(self.directory, "index.html")
            print(f"Serving index.html for {path}")
            if os.path.exists(index_path):
                return FileResponse(index_path)
        response = await super().get_response(path, scope)
        print(f"Response status for {path}: {response.status_code}")
        return response

# --- Expiry Mechanism ---
EXPIRY_FILE = os.path.join(os.path.dirname(__file__), "install_time.txt")
EXPIRY_DAYS = 3

def check_expiry():
    if not os.path.exists(EXPIRY_FILE):
        with open(EXPIRY_FILE, "w") as f:
            f.write(str(time.time()))
        return False  # Not expired
    with open(EXPIRY_FILE, "r") as f:
        first_time = float(f.read().strip())
    expiry_time = first_time + EXPIRY_DAYS * 24 * 60 * 60
    if time.time() > expiry_time:
        return True  # Expired
    return False  # Not expired

# --- Per-user trial remaining endpoint ---
from fastapi import Depends
@app.get("/trial-remaining")
def trial_remaining(current_user: User = Depends(get_current_user)):
    TRIAL_DAYS = 3
    if not current_user.created_at:
        return {"remaining_seconds": 0}
    expiry_time = current_user.created_at + datetime.timedelta(days=TRIAL_DAYS)
    remaining = int((expiry_time - datetime.datetime.utcnow()).total_seconds())
    if remaining < 0:
        remaining = 0
    return {"remaining_seconds": remaining}

if os.path.exists(frontend_build_dir):
    app.mount("/", SPAStaticFiles(directory=frontend_build_dir, html=True), name="static")
else:
    print(f"âš ï¸  Frontend build directory not found: {frontend_build_dir}")

REQUIRED_MODELS = [
    "llama3.1:latest",
    "mistral:latest",
    "anindya/prem1b-sql-ollama-fp116:latest",
    "deepseek-coder-v2:latest"  # Added for data analysis
]

def ensure_models():
    for model in REQUIRED_MODELS:
        try:
            print(f"Checking/pulling model: {model}")
            subprocess.run(["ollama", "pull", model], check=True)
        except Exception as e:
            print(f"Failed to pull model {model}: {e}")

# Ensure models are present before starting the server
ensure_models()

if __name__ == "__main__":
    import uvicorn
    # Use Render's PORT environment variable, fallback to 8000 for local development
    port = int(os.getenv("PORT", 8000))
    
    # Only open browser if not running from desktop launcher
    if not os.getenv("DESKTOP_MODE"):
        import threading
        import webbrowser
        def open_browser():
            webbrowser.open_new("http://localhost:8000")
        threading.Timer(1.5, open_browser).start()
    
    print(f"ðŸš€ Starting server on port {port}")
    print(f"ðŸŒ Environment: PORT={os.getenv('PORT', 'Not set (using 8000)')}")
    uvicorn.run(app, host="0.0.0.0", port=port)
